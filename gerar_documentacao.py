"""
Gera documentacao completa do projeto Brasileirao Dataset em formato DOCX.
Execute: python gerar_documentacao.py
Saida:   Brasileirao_Dataset_Documentacao.docx (raiz do projeto)
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paleta de cores ────────────────────────────────────────────────────────────
COR_TITULO      = RGBColor(0x1a, 0x37, 0x6d)
COR_SECAO       = RGBColor(0x1e, 0x5c, 0xb8)
COR_SUBSECAO    = RGBColor(0x2e, 0x7d, 0xd1)
COR_TEXTO       = RGBColor(0x1a, 0x1a, 0x2e)
COR_DESTAQUE    = RGBColor(0x00, 0x6e, 0x3c)
COR_AVISO       = RGBColor(0xc0, 0x39, 0x2b)
COR_BRANCO      = RGBColor(0xff, 0xff, 0xff)
FONTE_NORMAL    = "Calibri"
FONTE_CODIGO    = "Consolas"


# ── helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex6: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    tcPr.append(shd)


def _titulo_capa(doc, texto):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(26); run.font.bold = True
    run.font.name = FONTE_NORMAL; run.font.color.rgb = COR_TITULO


def _sub_capa(doc, texto):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(13); run.font.name = FONTE_NORMAL
    run.font.color.rgb = COR_SECAO


def h1(doc, texto):
    p = doc.add_heading(level=1); p.clear()
    r = p.add_run(texto)
    r.font.size = Pt(15); r.font.bold = True
    r.font.name = FONTE_NORMAL; r.font.color.rgb = COR_TITULO


def h2(doc, texto):
    p = doc.add_heading(level=2); p.clear()
    r = p.add_run(texto)
    r.font.size = Pt(12); r.font.bold = True
    r.font.name = FONTE_NORMAL; r.font.color.rgb = COR_SECAO


def h3(doc, texto):
    p = doc.add_heading(level=3); p.clear()
    r = p.add_run(texto)
    r.font.size = Pt(11); r.font.bold = True
    r.font.name = FONTE_NORMAL; r.font.color.rgb = COR_SUBSECAO


def para(doc, texto, negrito=False, italico=False, cor=None):
    p  = doc.add_paragraph()
    r  = p.add_run(texto)
    r.font.size = Pt(10.5); r.font.bold = negrito; r.font.italic = italico
    r.font.name = FONTE_NORMAL; r.font.color.rgb = cor or COR_TEXTO


def bul(doc, texto, nivel=0):
    p = doc.add_paragraph(style="List Bullet" if nivel == 0 else "List Bullet 2")
    r = p.add_run(texto)
    r.font.size = Pt(10.5); r.font.name = FONTE_NORMAL
    r.font.color.rgb = COR_TEXTO


def cod(doc, texto):
    """Bloco de codigo — fundo cinza claro, fonte monoespaco."""
    for linha in texto.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "ECF0F1")
        pPr.append(shd)
        r = p.add_run(linha if linha else " ")
        r.font.name = FONTE_CODIGO; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x27, 0x27, 0x27)


def aviso(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Atencao: " + texto)
    r.font.size = Pt(10); r.font.italic = True
    r.font.name = FONTE_NORMAL; r.font.color.rgb = COR_AVISO


def linha(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "1E5CB8")
    pBdr.append(bot); pPr.append(pBdr)


def tabela(doc, cabecalho, linhas, larguras=None):
    """Tabela com cabecalho azul e linhas alternadas."""
    t = doc.add_table(rows=1 + len(linhas), cols=len(cabecalho))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, txt in enumerate(cabecalho):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.font.bold = True; r.font.color.rgb = COR_BRANCO
        r.font.size = Pt(9.5); r.font.name = FONTE_NORMAL
        _set_cell_bg(c, "1A376D")

    for ri, row_data in enumerate(linhas):
        bg = "F2F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, txt in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9.5); r.font.name = FONTE_NORMAL
            r.font.color.rgb = COR_TEXTO
            _set_cell_bg(c, bg)

    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


# ── DOCUMENTO PRINCIPAL ────────────────────────────────────────────────────────

def build():
    doc = Document()

    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(2.0)

    # ============================================================
    # CAPA
    # ============================================================
    doc.add_paragraph()
    doc.add_paragraph()
    _titulo_capa(doc, "Brasileirao Data Lake")
    doc.add_paragraph()
    _sub_capa(doc, "Documentacao Tecnica e Funcional do Projeto")
    _sub_capa(doc, "Campeonato Brasileiro Serie A  -  2003 a 2026")
    doc.add_paragraph()
    linha(doc)
    doc.add_paragraph()

    for chave, valor in [
        ("Versao",      "1.0"),
        ("Data",        "Junho de 2026"),
        ("Autor",       "Lee O. Fernandes"),
        ("E-mail",      "leeofernandes@gmail.com"),
        ("Plataforma",  "Windows 11  Python 3.13  Node.js 22"),
        ("Repositorio", "Brasileirao_Dataset-master"),
    ]:
        p  = doc.add_paragraph()
        r1 = p.add_run(f"{chave}:  ")
        r1.font.bold = True; r1.font.size = Pt(10.5)
        r1.font.name = FONTE_NORMAL; r1.font.color.rgb = COR_TITULO
        r2 = p.add_run(valor)
        r2.font.size = Pt(10.5); r2.font.name = FONTE_NORMAL
        r2.font.color.rgb = COR_TEXTO

    doc.add_page_break()

    # ============================================================
    # 1. VISAO GERAL
    # ============================================================
    h1(doc, "1. Visao Geral do Projeto")
    linha(doc)
    para(doc,
        "O Brasileirao Data Lake e um projeto de engenharia de dados e visualizacao "
        "que consolida 24 edicoes do Campeonato Brasileiro Serie A (2003-2026) em um "
        "data lake estruturado em tres camadas (Bronze, Silver, Gold) e expoe os dados "
        "em um dashboard web interativo com tema dark inspirado no Sofascore.")
    doc.add_paragraph()

    h2(doc, "1.1 Objetivos")
    bul(doc, "Consolidar dados historicos (CSV 2003-2023) e dados em tempo real (API Sofascore 2024-2026) em formato padronizado Apache Parquet.")
    bul(doc, "Oferecer um pipeline ETL reprodutivel e automatizavel para atualizacao semanal via script .bat.")
    bul(doc, "Disponibilizar dashboard web moderno com analises interativas, incluindo confrontos diretos (H2H) com modelo probabilistico.")
    bul(doc, "Permitir analises ad-hoc via SQL interativo no browser com DuckDB-WASM.")
    doc.add_paragraph()

    h2(doc, "1.2 Tecnologias Utilizadas")
    tabela(doc,
        ["Camada",        "Tecnologia",            "Versao",   "Finalidade"],
        [
          ["Data Lake",   "Python",                "3.13",     "Pipeline ETL completo"],
          ["Data Lake",   "Pandas",                "2.1+",     "Transformacoes de dados"],
          ["Data Lake",   "PyArrow",               "15+",      "Serializacao Parquet"],
          ["Data Lake",   "DuckDB",                "0.10+",    "Consultas analiticas SQL"],
          ["Data Lake",   "python-dotenv",         "1.0+",     "Configuracao via .env"],
          ["API",         "Sofascore (RapidAPI)",  "-",        "Dados 2024-2026"],
          ["Dashboard",   "Next.js",               "16.2.7",   "Framework React SSR/CSR"],
          ["Dashboard",   "React",                 "19.2",     "Interface de usuario"],
          ["Dashboard",   "Tailwind CSS",          "4.x",      "Estilizacao responsiva"],
          ["Dashboard",   "Recharts",              "3.8",      "Graficos interativos"],
          ["Dashboard",   "DuckDB-WASM",           "1.33",     "SQL no browser"],
          ["Dashboard",   "Lucide React",          "1.17",     "Icones SVG"],
          ["Banco Rel.",  "MySQL",                 "8.0",      "Carga relacional (opcional)"],
        ],
        larguras=[1.2, 1.7, 0.8, 2.7]
    )

    doc.add_page_break()

    # ============================================================
    # 2. ARQUITETURA
    # ============================================================
    h1(doc, "2. Arquitetura do Sistema")
    linha(doc)
    para(doc, "O sistema e organizado em dois blocos: Data Lake (back-end) e Dashboard (front-end).")
    doc.add_paragraph()

    cod(doc,
"""FONTES DE DADOS
  CSV historico 2003-2023       -->  Bronze/csv/
  API Sofascore 2024-2026       -->  Bronze/api_cache/
               |
               |  Pipeline ETL (Python)
               v
SILVER LAYER  (Parquet padronizado)
  partidas . gols . cartoes . estatisticas . escalacoes
               |
               |  build_gold.py (DuckDB)
               v
GOLD LAYER  (Parquet analitico)
  campeonato_historico . artilharia . desempenho_clubes
  fair_play . classificacao_historica . rebaixamento_acesso
               |
               |  export_data.py
               v
DASHBOARD  (Next.js 16 / DuckDB-WASM)
  public/data/json/   <- JSONs pre-computados
  public/data/parquet/ <- Parquets p/ DuckDB no browser""")

    doc.add_paragraph()

    h2(doc, "2.1 Estrutura de Diretorios")
    cod(doc,
"""Brasileirao_Dataset-master/
|-- datalake/
|   |-- bronze/
|   |   |-- csv/              <- CSVs originais 2003-2023
|   |   |-- json/             <- JSONs por temporada
|   |   +-- api_cache/        <- Cache Sofascore (todas as temporadas)
|   |-- silver/
|   |   |-- partidas.parquet
|   |   |-- gols.parquet
|   |   |-- cartoes.parquet
|   |   |-- estatisticas.parquet
|   |   +-- escalacoes.parquet
|   |-- gold/
|   |   |-- campeonato_historico.parquet
|   |   |-- artilharia_historica.parquet
|   |   |-- desempenho_clubes.parquet
|   |   |-- fair_play.parquet
|   |   |-- classificacao_historica.parquet
|   |   +-- rebaixamento_acesso.parquet
|   |-- pipelines/
|   |   |-- config.py
|   |   |-- api_sofascore.py
|   |   |-- transform_silver.py
|   |   |-- ingest_sofascore.py
|   |   |-- ingest_historico_gols.py
|   |   |-- build_gold.py
|   |   |-- load_mysql.py
|   |   +-- run_all.py
|   |-- catalog/
|   |   +-- schema.yaml
|   +-- requirements.txt
|-- dashboard/
|   |-- src/app/
|   |   |-- page.tsx              <- Home
|   |   |-- temporadas/page.tsx   <- Lista de edicoes
|   |   |-- temporadas/[ano]/     <- Detalhes por temporada
|   |   |-- times/page.tsx        <- Grid de clubes
|   |   |-- times/[clube]/        <- Perfil do clube
|   |   |-- confrontos/page.tsx   <- H2H com probabilidade
|   |   +-- consultas/page.tsx    <- SQL interativo
|   |-- public/data/
|   |   |-- json/                 <- JSONs exportados
|   |   +-- parquet/              <- Parquets para DuckDB-WASM
|   +-- scripts/
|       +-- export_data.py
|-- atualizar_rodada.bat          <- Automacao Windows
+-- README.md""")

    doc.add_page_break()

    # ============================================================
    # 3. DATA LAKE — CAMADAS
    # ============================================================
    h1(doc, "3. Data Lake")
    linha(doc)

    h2(doc, "3.1 Camada Bronze")
    para(doc,
        "A camada Bronze preserva os dados originais sem qualquer modificacao, "
        "funcionando como staging area permanente para rastreabilidade completa.")
    doc.add_paragraph()
    tabela(doc,
        ["Diretorio",           "Conteudo",                             "Periodo"],
        [
          ["bronze/csv/",        "5 CSVs originais do dataset Adao Duque", "2003-2023"],
          ["bronze/json/",       "JSONs por temporada",                    "2003-2023"],
          ["bronze/api_cache/",  "Respostas brutas da API Sofascore",      "2003-2026"],
        ],
        larguras=[1.8, 3.5, 1.2]
    )
    bul(doc, "campeonato-brasileiro-full.csv      -- 9.568 partidas com resultados completos")
    bul(doc, "campeonato-brasileiro-gols.csv      -- artilheiros, minutos e tipos de gol")
    bul(doc, "campeonato-brasileiro-cartoes.csv   -- cartoes amarelos e vermelhos por partida")
    bul(doc, "campeonato-brasileiro-estatisticas-full.csv  -- chutes, passes, posse, faltas")
    doc.add_paragraph()

    h2(doc, "3.2 Camada Silver")
    para(doc,
        "Dados limpos com esquema padronizado. O campo temporada e derivado da data "
        "da partida com tratamento especial para o campeonato 2020 (jogos de jan-fev/2021). "
        "Todos os arquivos usam Apache Parquet com compressao Snappy.")
    doc.add_paragraph()

    h2(doc, "3.3 Camada Gold")
    para(doc,
        "Tabelas pre-computadas geradas pelo build_gold.py via DuckDB SQL sobre os "
        "Parquets da camada Silver. Prontas para consumo em dashboards e notebooks.")
    doc.add_paragraph()
    tabela(doc,
        ["Arquivo Gold",                    "Descricao",                                                  "Granularidade"],
        [
          ["campeonato_historico.parquet",   "Todas as partidas 2003-2026 enriquecidas",                  "1 linha / partida"],
          ["artilharia_historica.parquet",   "Gols por jogador por temporada (exclui gols contra)",       "1 linha / jogador / temp."],
          ["desempenho_clubes.parquet",      "Pontos, vitorias, saldo de gols e aproveitamento",          "1 linha / clube / temp."],
          ["fair_play.parquet",              "Amarelos, vermelhos e score de fair play",                  "1 linha / clube / temp."],
          ["classificacao_historica.parquet","Tabela de classificacao final calculada",                    "1 linha / clube / temp."],
          ["rebaixamento_acesso.parquet",    "Situacao final: Campeo / Libertadores / Rebaixado",         "1 linha / clube / temp."],
        ],
        larguras=[2.2, 3.2, 1.8]
    )

    doc.add_page_break()

    # ============================================================
    # 4. CATALOGO DE DADOS
    # ============================================================
    h1(doc, "4. Catalogo de Dados (Dicionario de Campos)")
    linha(doc)

    h2(doc, "4.1 partidas.parquet")
    para(doc, "Uma linha por partida disputada. Cobertura: 2003-2026.")
    tabela(doc,
        ["Coluna",               "Tipo",   "Descricao"],
        [
          ["partida_id",          "int",    "ID unico da partida"],
          ["temporada",           "int",    "Ano do campeonato (2003-2026)"],
          ["rodada",              "int",    "Numero da rodada (1-38)"],
          ["data",                "date",   "Data da partida"],
          ["hora",                "string", "Horario de inicio (HH:MM)"],
          ["mandante",            "string", "Clube da casa"],
          ["visitante",           "string", "Clube visitante"],
          ["formacao_mandante",   "string", "Formacao tatica do mandante (ex: 4-3-3)"],
          ["formacao_visitante",  "string", "Formacao tatica do visitante"],
          ["tecnico_mandante",    "string", "Tecnico do mandante"],
          ["tecnico_visitante",   "string", "Tecnico do visitante"],
          ["vencedor",            "string", "Nome do vencedor ou '-' para empate"],
          ["arena",               "string", "Nome do estadio"],
          ["gols_mandante",       "int",    "Gols marcados pelo mandante"],
          ["gols_visitante",      "int",    "Gols marcados pelo visitante"],
          ["estado_mandante",     "string", "UF do clube mandante"],
          ["estado_visitante",    "string", "UF do clube visitante"],
          ["status",              "string", "FT=finalizado / NS=nao iniciado / PST=adiado"],
          ["fonte",               "string", "historico_csv | sofascore"],
        ],
        larguras=[1.9, 0.7, 4.1]
    )

    h2(doc, "4.2 gols.parquet")
    para(doc, "Um registro por gol marcado. Cobertura parcial 2003-2011, completo 2012-2026.")
    tabela(doc,
        ["Coluna",      "Tipo",   "Descricao"],
        [
          ["partida_id", "int",    "Referencia a partidas.partida_id"],
          ["temporada",  "int",    "Ano do campeonato"],
          ["rodada",     "int",    "Numero da rodada"],
          ["clube",      "string", "Clube que marcou o gol"],
          ["atleta",     "string", "Nome do artilheiro"],
          ["minuto",     "int",    "Minuto em que o gol foi marcado"],
          ["tipo_de_gol","string", "Normal | Penalty | Gol Contra | Falta"],
          ["fonte",      "string", "historico_csv | sofascore"],
        ],
        larguras=[1.9, 0.7, 4.1]
    )

    h2(doc, "4.3 cartoes.parquet")
    para(doc, "Um registro por cartao aplicado.")
    tabela(doc,
        ["Coluna",      "Tipo",   "Descricao"],
        [
          ["partida_id", "int",    "Referencia a partidas.partida_id"],
          ["temporada",  "int",    "Ano do campeonato"],
          ["rodada",     "int",    "Numero da rodada"],
          ["clube",      "string", "Clube do jogador punido"],
          ["cartao",     "string", "Amarelo | Vermelho"],
          ["atleta",     "string", "Nome do jogador"],
          ["num_camisa", "string", "Numero da camisa"],
          ["posicao",    "string", "Posicao em campo"],
          ["minuto",     "int",    "Minuto da punicao"],
          ["fonte",      "string", "historico_csv | sofascore"],
        ],
        larguras=[1.9, 0.7, 4.1]
    )

    h2(doc, "4.4 estatisticas.parquet")
    para(doc, "Duas linhas por partida (uma por time). Dados esparsos em 2003-2011.")
    tabela(doc,
        ["Coluna",           "Tipo",   "Descricao"],
        [
          ["partida_id",      "int",    "Referencia a partidas.partida_id"],
          ["temporada",       "int",    "Ano do campeonato"],
          ["rodada",          "int",    "Numero da rodada"],
          ["clube",           "string", "Clube referenciado"],
          ["chutes",          "int",    "Total de finalizacoes"],
          ["chutes_no_alvo",  "int",    "Finalizacoes no gol"],
          ["posse_de_bola",   "float",  "Posse de bola em porcentagem"],
          ["passes",          "int",    "Total de passes"],
          ["precisao_passes", "float",  "Precisao de passe em porcentagem"],
          ["faltas",          "int",    "Faltas cometidas"],
          ["cartao_amarelo",  "int",    "Cartoes amarelos recebidos"],
          ["cartao_vermelho", "int",    "Cartoes vermelhos recebidos"],
          ["impedimentos",    "int",    "Impedimentos marcados"],
          ["escanteios",      "int",    "Escanteios cobrados"],
          ["fonte",           "string", "historico_csv | sofascore"],
        ],
        larguras=[1.9, 0.7, 4.1]
    )

    h2(doc, "4.5 escalacoes.parquet")
    para(doc, "Elencos em campo. Disponivel apenas para 2024+ (via API Sofascore).")
    tabela(doc,
        ["Coluna",      "Tipo",   "Descricao"],
        [
          ["partida_id", "int",    "Referencia a partidas.partida_id"],
          ["temporada",  "int",    "Ano do campeonato"],
          ["rodada",     "int",    "Numero da rodada"],
          ["clube",      "string", "Clube"],
          ["tecnico",    "string", "Tecnico responsavel"],
          ["formacao",   "string", "Formacao tatica (ex: 4-3-3)"],
          ["atleta",     "string", "Nome do jogador"],
          ["num_camisa", "int",    "Numero da camisa"],
          ["posicao",    "string", "Posicao em campo"],
          ["titular",    "bool",   "True = titular  /  False = reserva que entrou"],
          ["fonte",      "string", "sofascore"],
        ],
        larguras=[1.9, 0.7, 4.1]
    )

    doc.add_page_break()

    # ============================================================
    # 5. PIPELINE ETL
    # ============================================================
    h1(doc, "5. Pipeline ETL")
    linha(doc)
    para(doc,
        "O pipeline e composto por scripts Python independentes em datalake/pipelines/. "
        "Cada script tem responsabilidade unica e pode ser executado de forma isolada.")
    doc.add_paragraph()

    h2(doc, "5.1 Scripts e Responsabilidades")
    tabela(doc,
        ["Script",                     "Responsabilidade",                                          "Requer API"],
        [
          ["transform_silver.py",       "Converte CSVs historicos 2003-2023 em Parquet Silver",      "Nao"],
          ["ingest_sofascore.py",       "Ingere dados 2024-2026 via API Sofascore (upsert)",         "Sim"],
          ["ingest_historico_gols.py",  "Backfill de gols e cartoes 2003-2013 via API",              "Sim"],
          ["build_gold.py",             "Gera tabelas analiticas Gold via DuckDB SQL",               "Nao"],
          ["export_data.py",            "Exporta JSON e Parquet para public/data/ do dashboard",     "Nao"],
          ["load_mysql.py",             "Carrega tabelas Gold no banco MySQL (opcional)",             "Nao"],
          ["run_all.py",                "Orquestra todos os scripts na ordem correta",               "Opcional"],
          ["update_season.py",          "Atualizacao rapida de uma temporada especifica",            "Sim"],
          ["config.py",                 "Caminhos e constantes do projeto",                          "Nao"],
          ["api_sofascore.py",          "Cliente da API Sofascore com cache e normalizacao de nomes","Sim"],
        ],
        larguras=[2.2, 3.8, 0.8]
    )

    h2(doc, "5.2 Ordem de Execucao Recomendada")
    tabela(doc,
        ["Etapa", "Comando",                                                         "Descricao"],
        [
          ["1",    "python transform_silver.py",                                     "CSV historico 2003-2023 --> Parquet"],
          ["2",    "python ingest_sofascore.py --seasons 2024 2025 2026",            "Upsert dados 2024-2026"],
          ["3",    "python ingest_historico_gols.py",                               "Backfill gols/cartoes 2003-2013"],
          ["4",    "python build_gold.py",                                           "Gera tabelas Gold"],
          ["5",    "python export_data.py",                                          "Exporta para dashboard"],
          ["6",    "python load_mysql.py",                                           "Carga MySQL (opcional)"],
        ],
        larguras=[0.5, 3.5, 2.8]
    )
    aviso(doc,
        "transform_silver.py SOBRESCREVE partidas.parquet por completo. "
        "Execute SEMPRE ingest_sofascore.py apos ele para restaurar os dados 2024-2026.")
    doc.add_paragraph()

    h2(doc, "5.3 Comandos de Execucao")
    para(doc, "Primeira execucao completa sem API (apenas historico):", negrito=True)
    cod(doc,
"""cd datalake/pipelines
python run_all.py --skip-api""")

    para(doc, "Primeira execucao completa com API:", negrito=True)
    cod(doc,
"""cd datalake/pipelines
python run_all.py --fixtures-only""")

    para(doc, "Ingestao por temporada especifica:", negrito=True)
    cod(doc,
"""python ingest_sofascore.py --seasons 2026 --fixtures-only   # apenas placares
python ingest_sofascore.py --seasons 2026                     # placares + gols + cartoes""")

    para(doc, "Atualizacao semanal rapida:", negrito=True)
    cod(doc,
"""python update_season.py --fixtures-only --rebuild-gold""")
    doc.add_paragraph()

    h2(doc, "5.4 Cotas da API Sofascore")
    tabela(doc,
        ["Modo",           "Requisicoes / Temporada", "Observacao"],
        [
          ["fixtures-only", "~39",                     "1 (rounds) + 38 (partidas). Cabe no free tier de 100/dia."],
          ["Completo",      "~1.179",                  "39 + 3 req x 380 partidas (stats, gols, cartoes)"],
          ["Free Tier",     "100 / dia",               "Suficiente para fixtures-only semanal"],
        ],
        larguras=[1.5, 2.0, 3.5]
    )

    h2(doc, "5.5 Tratamento do COVID-19 (Brasileirao 2020)")
    para(doc,
        "As rodadas 28-38 do Brasileirao 2020 foram realizadas em jan-fev/2021 devido "
        "a pandemia. O pipeline detecta automaticamente partidas com "
        "year==2021 AND month<=3 e as reclassifica como temporada 2020.")
    doc.add_paragraph()

    h2(doc, "5.6 Normalizacao de Nomes de Clubes")
    para(doc,
        "O modulo api_sofascore.py mantem o dicionario _TEAM_NAMES que mapeia nomes "
        "da API Sofascore para os nomes canonicos do CSV historico. Exemplos:")
    bul(doc, "'Atletico Mineiro'    -->  'Atletico-MG'")
    bul(doc, "'Red Bull Bragantino' -->  'Bragantino'")
    bul(doc, "'Athletico Paranaense'-->  'Atletico-PR'")

    doc.add_page_break()

    # ============================================================
    # 6. AUTOMACAO
    # ============================================================
    h1(doc, "6. Automacao Windows — atualizar_rodada.bat")
    linha(doc)
    para(doc,
        "O arquivo atualizar_rodada.bat na raiz do projeto atualiza todo o stack "
        "com um duplo clique. Detecta automaticamente o Python instalado no sistema.")
    doc.add_paragraph()

    tabela(doc,
        ["Etapa", "Acao",                                           "Script"],
        [
          ["1/5",  "Busca placares da ultima rodada",               "ingest_sofascore.py --seasons 2026 --fixtures-only"],
          ["2/5",  "Busca gols, cartoes e escalacoes",              "ingest_sofascore.py --seasons 2026"],
          ["3/5",  "Reconstroi tabelas analiticas Gold",            "build_gold.py"],
          ["4/5",  "Exporta dados para o dashboard",                "export_data.py"],
          ["5/5",  "Atualiza banco MySQL",                          "load_mysql.py"],
        ],
        larguras=[0.5, 3.0, 3.5]
    )
    aviso(doc,
        "Salvar o .bat com CRLF e codificacao UTF-8 sem BOM. "
        "Arquivos com terminadores LF-only fragmentam comandos no CMD do Windows.")
    doc.add_paragraph()

    para(doc, "Deteccao automatica do Python:", negrito=True)
    cod(doc,
"""set PYTHON=python
if exist "C:\\Python314\\python.exe" set PYTHON=C:\\Python314\\python.exe
if exist "C:\\Python313\\python.exe" set PYTHON=C:\\Python313\\python.exe
if exist "C:\\Python312\\python.exe" set PYTHON=C:\\Python312\\python.exe""")

    doc.add_page_break()

    # ============================================================
    # 7. DASHBOARD WEB
    # ============================================================
    h1(doc, "7. Dashboard Web")
    linha(doc)
    para(doc,
        "Aplicacao Next.js 16 com React 19 e Tailwind CSS 4, tema dark inspirado "
        "no Sofascore. Dados carregados via fetch de JSONs estaticos; pagina de consultas "
        "usa DuckDB-WASM para SQL direto no browser sem servidor.")
    doc.add_paragraph()

    h2(doc, "7.1 Tema Visual (globals.css)")
    tabela(doc,
        ["Variavel CSS",   "Valor",    "Uso"],
        [
          ["--bg",          "#1a1d29", "Fundo principal da pagina"],
          ["--bg-card",     "#222636", "Fundo de cards e tabelas"],
          ["--border",      "#2d3250", "Bordas e divisorias"],
          ["--nav",         "#1e2235", "Fundo da barra de navegacao"],
          ["--fg",          "#e8eaf6", "Texto principal"],
          ["--muted",       "#8b92b3", "Texto secundario e labels"],
          ["--accent",      "#3d7cf5", "Cor de destaque (azul)"],
          ["--green",       "#00d25b", "Vitorias e resultados positivos"],
          ["--yellow",      "#f5c518", "Empates e alertas"],
          ["--red",         "#e94560", "Derrotas e alertas negativos"],
        ],
        larguras=[1.5, 1.1, 4.0]
    )

    h2(doc, "7.2 Paginas e Funcionalidades")
    tabela(doc,
        ["Rota",               "Arquivo",                             "Funcionalidades Principais"],
        [
          ["/",                 "app/page.tsx",                       "KPIs totais, grafico gols/temporada, ranking de campeoes"],
          ["/temporadas",       "app/temporadas/page.tsx",            "Lista todas as edicoes 2003-2026 com estatisticas resumidas"],
          ["/temporadas/[ano]", "app/temporadas/[ano]/page.tsx",      "Classificacao, artilharia e lista de rodadas por ano"],
          ["/times",            "app/times/page.tsx",                 "Grid de todos os clubes com avatares coloridos gerados"],
          ["/times/[clube]",    "app/times/[clube]/page.tsx",         "Perfil: grafico de aproveitamento e historico de posicoes"],
          ["/confrontos",       "app/confrontos/page.tsx",            "H2H com modelo probabilistico, forma recente e historico"],
          ["/consultas",        "app/consultas/page.tsx",             "SQL interativo com DuckDB-WASM no browser"],
        ],
        larguras=[1.5, 2.4, 3.0]
    )

    h2(doc, "7.3 KPIs da Home Page")
    bul(doc, "Total de partidas realizadas no periodo 2003-2026")
    bul(doc, "Total de gols marcados no periodo")
    bul(doc, "Media de gols por partida")
    bul(doc, "Total de temporadas cobertas")
    bul(doc, "Grafico de barras: gols totais por temporada")
    bul(doc, "Ranking de campeoes: titulos por clube")
    doc.add_paragraph()

    h2(doc, "7.4 Modelo Probabilistico (Pagina Confrontos)")
    para(doc,
        "O modelo combina tres fatores com pesos configurados pelo modo de elenco escolhido:")
    doc.add_paragraph()
    bul(doc, "Fator 1 -- H2H com Decaimento Temporal (DECAY = 0,82 por ano): "
             "peso de um jogo de 10 anos atras = ~14% de um jogo recente.")
    bul(doc, "Fator 2 -- Forma Recente (ultimos N jogos da Serie A): "
             "janela configuravel por modo de elenco.")
    bul(doc, "Fator 3 -- Forca do Elenco (peso fixo 25%): pontos por jogo na ultima "
             "temporada completa, com penalidade para times fora da Serie A.")
    doc.add_paragraph()

    tabela(doc,
        ["Modo Elenco", "Janela Forma", "Peso H2H", "Descricao"],
        [
          ["Titular",    "5 jogos",      "55%",      "Confia mais na forma atual"],
          ["Misto",      "8 jogos",      "70%",      "Balanco entre historico e forma recente"],
          ["Reserva",    "12 jogos",     "82%",      "Prioriza historico e H2H"],
        ],
        larguras=[1.2, 1.2, 1.0, 4.0]
    )
    bul(doc, "Controle de mando de campo: Casa / Fora / Neutro")
    bul(doc, "Times ausentes da Serie A ha 2+ anos recebem aviso e penalidade na forca")
    bul(doc, "Tabs: Retrospecto / Probabilidade / Forma / Historico de partidas")
    doc.add_paragraph()

    h2(doc, "7.5 SQL Interativo (DuckDB-WASM)")
    para(doc,
        "A pagina /consultas carrega os Parquets de public/data/parquet/ no browser "
        "via DuckDB-WASM, sem necessidade de servidor back-end. O usuario executa "
        "qualquer SELECT SQL sobre as tabelas do data lake.")
    cod(doc,
"""-- Top 10 artilheiros historicos (excluindo gols contra)
SELECT atleta, clube, SUM(total_gols) AS gols
FROM artilharia_historica
GROUP BY atleta, clube
ORDER BY gols DESC
LIMIT 10;

-- Historico do Flamengo na Serie A
SELECT temporada, pontos, vitorias, saldo_gols, aproveitamento_pct
FROM desempenho_clubes
WHERE clube = 'Flamengo'
ORDER BY temporada;

-- Todos os rebaixamentos
SELECT temporada, clube, posicao, pontos
FROM rebaixamento_acesso
WHERE situacao = 'Rebaixado'
ORDER BY temporada;""")

    doc.add_page_break()

    # ============================================================
    # 8. DADOS EXPORTADOS
    # ============================================================
    h1(doc, "8. Dados Exportados para o Dashboard")
    linha(doc)
    para(doc,
        "O script dashboard/scripts/export_data.py gera os arquivos estaticos "
        "consumidos pelo dashboard em public/data/.")
    doc.add_paragraph()

    tabela(doc,
        ["Arquivo JSON",                  "Conteudo",                                   "Volume (aprox.)"],
        [
          ["partidas.json",               "Todas as partidas 2003-2026",                "~2,5 MB"],
          ["gols.json",                   "Todos os gols registrados (11.285+)",        "~1,2 MB"],
          ["campeonato_historico.json",   "Partidas Gold enriquecidas",                 "~3,0 MB"],
          ["artilharia_historica.json",   "Artilharia por jogador e temporada",         "~0,8 MB"],
          ["desempenho_clubes.json",      "Desempenho por clube e temporada",           "~0,2 MB"],
          ["classificacao_historica.json","Classificacao final de cada temporada",      "~0,3 MB"],
          ["rebaixamento_acesso.json",    "Situacao final de cada clube",               "~0,1 MB"],
          ["fair_play.json",              "Cartoes e score de fair play por clube",     "~0,1 MB"],
          ["temporadas_meta.json",        "Metadados resumidos de cada temporada",      "~0,05 MB"],
        ],
        larguras=[2.5, 3.0, 1.8]
    )
    aviso(doc,
        "Antes do deploy na Vercel, verificar se public/data/ ultrapassa o limite de "
        "100 MB. Se sim, usar Vercel Blob Storage, AWS S3 ou Cloudflare R2 como CDN externo.")

    doc.add_page_break()

    # ============================================================
    # 9. COBERTURA
    # ============================================================
    h1(doc, "9. Cobertura dos Dados por Periodo")
    linha(doc)
    tabela(doc,
        ["Tabela",               "2003-2011",      "2012-2023",   "2024-2026"],
        [
          ["Partidas (results.)", "Completo",        "Completo",    "Via API Sofascore"],
          ["Estatisticas",        "Parcial/Esparso", "Completo",    "Via API Sofascore"],
          ["Gols (artilharia)",   "Parcial *",       "Completo",    "Via API Sofascore"],
          ["Cartoes",             "Parcial *",       "Completo",    "Via API Sofascore"],
          ["Escalacoes",          "Nao disponivel",  "Nao disp.",   "Via API Sofascore"],
          ["Classificacao",       "Calculada",       "Calculada",   "Via API Sofascore"],
        ],
        larguras=[1.8, 1.8, 1.5, 2.0]
    )
    para(doc,
        "* Backfill de gols e cartoes 2003-2013 implementado via ingest_historico_gols.py, "
        "que usa a API publica do Sofascore com cache local em bronze/api_cache/. "
        "Os resultados ficam cacheados: re-executar e gratuito.")

    doc.add_page_break()

    # ============================================================
    # 10. INSTALACAO
    # ============================================================
    h1(doc, "10. Instalacao e Configuracao")
    linha(doc)

    h2(doc, "10.1 Pre-requisitos")
    bul(doc, "Python 3.12 ou superior (testado em 3.13)")
    bul(doc, "Node.js 22 + npm (para o dashboard Next.js)")
    bul(doc, "Conta RapidAPI com assinatura do Sofascore (free tier: 100 req/dia)")
    bul(doc, "MySQL 8.0 (opcional -- apenas para carga relacional)")
    doc.add_paragraph()

    h2(doc, "10.2 Instalacao do Data Lake")
    cod(doc,
"""# 1. Instalar dependencias Python
pip install -r datalake/requirements.txt

# 2. Configurar variaveis de ambiente
copy datalake\\.env.example datalake\\.env
# Editar .env e preencher SOFASCORE_API_KEY

# 3. Pipeline completo sem API (historico apenas)
cd datalake/pipelines
python run_all.py --skip-api

# 4. Pipeline com API (inclui 2024-2026)
python run_all.py --fixtures-only""")
    doc.add_paragraph()

    h2(doc, "10.3 Instalacao do Dashboard")
    cod(doc,
"""cd dashboard
npm install

# Exportar dados do datalake para o dashboard
python scripts/export_data.py

# Servidor de desenvolvimento (http://localhost:3000)
npm run dev

# Build de producao
npm run build && npm start""")
    doc.add_paragraph()

    h2(doc, "10.4 Variaveis de Ambiente (.env)")
    tabela(doc,
        ["Variavel",            "Obrigatoria", "Descricao"],
        [
          ["SOFASCORE_API_KEY", "Sim",         "Chave da API Sofascore no RapidAPI"],
          ["MYSQL_HOST",        "Nao",         "Host do servidor MySQL (default: localhost)"],
          ["MYSQL_PORT",        "Nao",         "Porta MySQL (default: 3306)"],
          ["MYSQL_USER",        "Nao",         "Usuario MySQL (default: root)"],
          ["MYSQL_PASSWORD",    "Nao",         "Senha do MySQL"],
          ["MYSQL_DATABASE",    "Nao",         "Nome do banco (default: brasileirao_datalake)"],
        ],
        larguras=[2.0, 1.2, 4.0]
    )

    doc.add_page_break()

    # ============================================================
    # 11. TROUBLESHOOTING
    # ============================================================
    h1(doc, "11. Problemas Conhecidos e Solucoes")
    linha(doc)

    h2(doc, "11.1 Partidas 2024-2026 desaparecem apos transform_silver.py")
    para(doc, "Causa: transform_silver.py sobrescreve partidas.parquet do zero a partir dos CSVs historicos (cobrem apenas ate 2023).")
    para(doc, "Solucao:", negrito=True)
    cod(doc, "python ingest_sofascore.py --seasons 2024 2025 2026")
    doc.add_paragraph()

    h2(doc, "11.2 Rotas 404 no Next.js 16 (routes.d.ts corrompido)")
    para(doc, "Causa: arquivo .next/dev/types/routes.d.ts pode se corromper em hot-reload.")
    para(doc, "Solucao:", negrito=True)
    cod(doc,
"""rd /s /q dashboard\\.next
cd dashboard && npm run dev""")
    doc.add_paragraph()

    h2(doc, "11.3 Arquivo .bat com terminadores LF-only")
    para(doc, "Causa: editores como VS Code podem salvar .bat com LF, fragmentando comandos no CMD.")
    para(doc, "Solucao: Salvar com CRLF + UTF-8 sem BOM. VS Code: canto inferior direito > alterar para CRLF.")
    doc.add_paragraph()

    h2(doc, "11.4 Cota da API esgotada (429)")
    para(doc,
        "Use --fixtures-only (39 req/temporada) em vez do modo completo (1.179 req). "
        "O cache em bronze/api_cache/ evita requisicoes duplicadas.")
    doc.add_paragraph()

    h2(doc, "11.5 Encoding errado no CMD Windows")
    para(doc, "Execute chcp 65001 no CMD para ativar UTF-8 (ja incluso no .bat).")

    doc.add_page_break()

    # ============================================================
    # 12. DEPLOY
    # ============================================================
    h1(doc, "12. Publicacao e Deploy")
    linha(doc)

    h2(doc, "12.1 GitHub")
    cod(doc,
"""git init
git add .
git commit -m "feat: Brasileirao Data Lake v1.0"
git remote add origin https://github.com/<usuario>/brasileirao-data-lake.git
git push -u origin main""")
    aviso(doc, "Nao versionar: datalake/.env   datalake/bronze/api_cache/   node_modules/")
    doc.add_paragraph()

    h2(doc, "12.2 Vercel (Dashboard)")
    cod(doc,
"""npm i -g vercel
cd dashboard
vercel --prod""")
    para(doc, "vercel.json recomendado:", negrito=True)
    cod(doc,
"""{
  "framework": "nextjs",
  "buildCommand": "npm run build",
  "outputDirectory": ".next",
  "installCommand": "npm install"
}""")
    aviso(doc,
        "Limite Vercel Free: 100 MB por deploy. Verificar tamanho de public/data/ "
        "antes do deploy. Se ultrapassar, usar Vercel Blob Storage ou CDN externo.")
    doc.add_paragraph()

    # ============================================================
    # 13. FONTES
    # ============================================================
    h1(doc, "13. Fontes de Dados e Creditos")
    linha(doc)
    tabela(doc,
        ["Fonte",                   "Periodo",    "Responsavel",          "Contato"],
        [
          ["Dataset CSV Brasileiro", "2003-2023",  "Adao Duque",           "adaoduquesn@gmail.com"],
          ["API Sofascore RapidAPI", "2003-2026",  "Sofascore / RapidAPI", "rapidapi.com"],
        ],
        larguras=[2.0, 1.0, 1.8, 2.8]
    )
    bul(doc, "Dataset original disponivel no Kaggle: campeonato-brasileiro-serie-a")
    bul(doc, "Liga ID Brasileirao Serie A no Sofascore = 325")
    bul(doc, "Free tier RapidAPI: 100 requisicoes/dia (suficiente para atualizacoes semanais)")
    doc.add_paragraph()

    # ============================================================
    # CAPA FINAL
    # ============================================================
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph()
    _titulo_capa(doc, "Fim da Documentacao")
    doc.add_paragraph()
    _sub_capa(doc, "Brasileirao Data Lake  --  Serie A 2003-2026")
    _sub_capa(doc, "Junho de 2026")
    doc.add_paragraph()
    linha(doc)
    doc.add_paragraph()
    para(doc,
        "Este documento foi gerado automaticamente pelo script gerar_documentacao.py "
        "localizado na raiz do projeto. Para regenerar apos atualizacoes, execute:",
        italico=True)
    cod(doc, "python gerar_documentacao.py")

    # ── salvar ─────────────────────────────────────────────────────────────────
    out = Path(__file__).parent / "Brasileirao_Dataset_Documentacao.docx"
    doc.save(str(out))
    print(f"\nDocumento salvo em: {out}")
    return str(out)


if __name__ == "__main__":
    build()
