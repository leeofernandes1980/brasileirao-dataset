"""Gera DOCX com queries SQL de treinamento para o banco brasileirao_datalake."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()

# ── Estilos globais ────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)


def titulo(texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)
    return h


def subtitulo(texto):
    return titulo(texto, nivel=2)


def descricao(texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(2)
    return p


def sql(codigo):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(8)
    run = p.add_run(codigo.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # fundo cinza claro via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def separador():
    doc.add_paragraph("─" * 80)


# ══════════════════════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BRASILEIRÃO DATA LAKE")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0x9C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guia de Consultas SQL para Treinamento")
r.bold = True
r.font.size = Pt(15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Banco de dados: brasileirao_datalake  |  MySQL 8.0")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Temporadas 2003 – 2026  |  9.568 partidas")
r.font.size = Pt(11)

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SUMÁRIO (manual)
# ══════════════════════════════════════════════════════════════════════════════
titulo("Estrutura do Banco de Dados")
descricao("Antes de iniciar, ative o banco com o comando abaixo:")
sql("USE brasileirao_datalake;")

descricao("Tabelas disponíveis:")
tabelas = [
    ("partidas",               "Todos os jogos (2003-2026): placar, data, times, rodada, status"),
    ("gols",                   "Gols por partida: atleta, minuto, tipo (normal, pênalti, contra)"),
    ("cartoes",                "Cartões amarelos e vermelhos por partida"),
    ("estatisticas",           "Estatísticas por partida: posse, chutes, escanteios, faltas"),
    ("campeonato_historico",   "Cópia completa de partidas — base para análises históricas"),
    ("classificacao_historica","Tabela de classificação final por temporada e clube"),
    ("artilharia_historica",   "Gols marcados por jogador agrupados por temporada"),
    ("desempenho_clubes",      "Aproveitamento histórico de cada clube"),
    ("fair_play",              "Total de cartões por clube e temporada"),
    ("rebaixamento_acesso",    "Posição final de cada clube por temporada"),
]
for nome, desc in tabelas:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(nome.ljust(30))
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    p.add_run(f" — {desc}").font.size = Pt(10)

doc.add_paragraph()
descricao("Views pré-construídas:")
views = [
    ("v_artilharia_geral",    "Top artilheiros de todos os tempos"),
    ("v_campeoes",            "Campeão de cada temporada"),
    ("v_rebaixados",          "Times rebaixados por temporada"),
    ("v_desempenho_geral",    "Aproveitamento histórico consolidado"),
    ("v_media_gols_temporada","Média de gols por jogo em cada temporada"),
]
for nome, desc in views:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(nome.ljust(28))
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    p.add_run(f" — {desc}").font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NÍVEL 1 — BÁSICO
# ══════════════════════════════════════════════════════════════════════════════
titulo("Nível 1 — Consultas Básicas")

subtitulo("1.1 Exploração inicial")
descricao("Total de partidas no banco:")
sql("SELECT COUNT(*) AS total_partidas FROM partidas;")

descricao("Total de partidas por temporada (mais recentes primeiro):")
sql("""\
SELECT temporada, COUNT(*) AS total
FROM partidas
GROUP BY temporada
ORDER BY temporada DESC;""")

descricao("Primeiras 10 partidas da rodada 1 de 2024:")
sql("""\
SELECT rodada, data, mandante, gols_mandante, gols_visitante, visitante, arena
FROM partidas
WHERE temporada = 2024 AND rodada = 1
ORDER BY data;""")

descricao("Partidas com mais gols na história:")
sql("""\
SELECT data, mandante, gols_mandante, gols_visitante, visitante,
       (gols_mandante + gols_visitante) AS total_gols
FROM partidas
WHERE gols_mandante IS NOT NULL
ORDER BY total_gols DESC
LIMIT 10;""")

subtitulo("1.2 Filtragem e condições")
descricao("Todos os jogos do Flamengo em 2024:")
sql("""\
SELECT data, rodada, mandante, gols_mandante, gols_visitante, visitante
FROM partidas
WHERE temporada = 2024
  AND (mandante = 'Flamengo' OR visitante = 'Flamengo')
ORDER BY data;""")

descricao("Clássicos São Paulo x Corinthians ao longo da história:")
sql("""\
SELECT temporada, data, rodada, gols_mandante, gols_visitante
FROM partidas
WHERE (mandante = 'São Paulo' AND visitante = 'Corinthians')
   OR (mandante = 'Corinthians' AND visitante = 'São Paulo')
ORDER BY data;""")

descricao("Partidas encerradas com resultado (apenas jogos finalizados):")
sql("""\
SELECT temporada, data, mandante, gols_mandante, gols_visitante, visitante
FROM partidas
WHERE status = 'finished'
ORDER BY data DESC
LIMIT 20;""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NÍVEL 2 — INTERMEDIÁRIO
# ══════════════════════════════════════════════════════════════════════════════
titulo("Nível 2 — Agregações e Agrupamentos")

subtitulo("2.1 Classificação e pontuação")
descricao("Classificação completa de 2024 (com pontos calculados):")
sql("""\
SELECT
    ROW_NUMBER() OVER (ORDER BY pontos DESC, vitorias DESC) AS pos,
    clube,
    pontos,
    vitorias AS V,
    empates  AS E,
    derrotas AS D,
    gols_pro   AS GP,
    gols_contra AS GC,
    (gols_pro - gols_contra) AS SG
FROM classificacao_historica
WHERE temporada = 2024
ORDER BY pontos DESC, vitorias DESC;""")

descricao("Classificação parcial de 2026 (temporada em andamento):")
sql("""\
SELECT
    ROW_NUMBER() OVER (ORDER BY pontos DESC, vitorias DESC) AS pos,
    clube, pontos, vitorias AS V, empates AS E, derrotas AS D
FROM classificacao_historica
WHERE temporada = 2026
ORDER BY pontos DESC, vitorias DESC;""")

descricao("Campeões históricos (usando a view pronta):")
sql("SELECT * FROM v_campeoes ORDER BY temporada DESC;")

descricao("Quantas vezes cada clube foi campeão:")
sql("""\
SELECT clube, COUNT(*) AS titulos
FROM v_campeoes
GROUP BY clube
ORDER BY titulos DESC;""")

subtitulo("2.2 Análise de gols")
descricao("Média de gols por jogo em cada temporada:")
sql("SELECT * FROM v_media_gols_temporada ORDER BY temporada DESC;""")

descricao("Temporada com mais gols no total:")
sql("""\
SELECT temporada,
       SUM(gols_mandante + gols_visitante) AS total_gols,
       COUNT(*) AS partidas,
       ROUND(SUM(gols_mandante + gols_visitante) / COUNT(*), 2) AS media
FROM partidas
WHERE gols_mandante IS NOT NULL
GROUP BY temporada
ORDER BY total_gols DESC;""")

descricao("Artilharia histórica — top 20 goleadores:")
sql("SELECT * FROM v_artilharia_geral LIMIT 20;")

descricao("Artilharia de 2024 (temporada completa):")
sql("""\
SELECT atleta, clube, SUM(1) AS gols, temporada
FROM gols
WHERE temporada = 2024
GROUP BY atleta, clube, temporada
ORDER BY gols DESC
LIMIT 20;""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NÍVEL 3 — AVANÇADO
# ══════════════════════════════════════════════════════════════════════════════
titulo("Nível 3 — Análises Avançadas")

subtitulo("3.1 Desempenho histórico de clubes")
descricao("Aproveitamento histórico de todos os clubes (total de pontos e % aproveitamento):")
sql("""\
SELECT
    clube,
    SUM(pontos)   AS pontos_total,
    SUM(vitorias) AS vitorias,
    SUM(empates)  AS empates,
    SUM(derrotas) AS derrotas,
    COUNT(*)      AS temporadas,
    ROUND(SUM(pontos) / (COUNT(*) * 114) * 100, 1) AS aproveitamento_pct
FROM classificacao_historica
GROUP BY clube
ORDER BY pontos_total DESC
LIMIT 20;""")

descricao("Evolução do Flamengo temporada a temporada:")
sql("""\
SELECT temporada, pontos, vitorias, empates, derrotas,
       gols_pro, gols_contra,
       (gols_pro - gols_contra) AS saldo
FROM classificacao_historica
WHERE clube = 'Flamengo'
ORDER BY temporada;""")

descricao("Times que mais sofreram rebaixamentos:")
sql("""\
SELECT clube, COUNT(*) AS rebaixamentos
FROM v_rebaixados
GROUP BY clube
ORDER BY rebaixamentos DESC
LIMIT 15;""")

subtitulo("3.2 Análise de partidas em casa vs. fora")
descricao("Aproveitamento mandante vs. visitante por temporada:")
sql("""\
SELECT
    temporada,
    COUNT(*) AS partidas,
    SUM(CASE WHEN gols_mandante > gols_visitante THEN 1 ELSE 0 END) AS vit_mandante,
    SUM(CASE WHEN gols_mandante = gols_visitante THEN 1 ELSE 0 END) AS empates,
    SUM(CASE WHEN gols_mandante < gols_visitante THEN 1 ELSE 0 END) AS vit_visitante,
    ROUND(SUM(CASE WHEN gols_mandante > gols_visitante THEN 1 ELSE 0 END)
          / COUNT(*) * 100, 1) AS pct_casa
FROM partidas
WHERE gols_mandante IS NOT NULL
GROUP BY temporada
ORDER BY temporada DESC;""")

descricao("Clubes com melhor aproveitamento como visitante (histórico):")
sql("""\
SELECT
    visitante AS clube,
    COUNT(*) AS jogos_fora,
    SUM(CASE WHEN gols_visitante > gols_mandante THEN 1 ELSE 0 END) AS vitorias_fora,
    ROUND(SUM(CASE WHEN gols_visitante > gols_mandante THEN 1 ELSE 0 END)
          / COUNT(*) * 100, 1) AS pct_vitoria_fora
FROM partidas
WHERE gols_mandante IS NOT NULL AND status = 'finished'
GROUP BY visitante
HAVING jogos_fora >= 50
ORDER BY pct_vitoria_fora DESC
LIMIT 15;""")

doc.add_page_break()

subtitulo("3.3 Estatísticas de partidas")
descricao("Médias de estatísticas por temporada (posse, chutes, escanteios):")
sql("""\
SELECT
    p.temporada,
    ROUND(AVG(e.posse_de_bola), 1)    AS media_posse,
    ROUND(AVG(e.chutes), 1)           AS media_chutes,
    ROUND(AVG(e.escanteios), 1)       AS media_escanteios,
    ROUND(AVG(e.faltas), 1)           AS media_faltas
FROM estatisticas e
JOIN partidas p ON e.partida_id = p.partida_id
GROUP BY p.temporada
ORDER BY p.temporada DESC;""")

descricao("Partidas com mais cartões:")
sql("""\
SELECT
    p.data, p.mandante, p.gols_mandante, p.gols_visitante, p.visitante,
    COUNT(c.partida_id) AS total_cartoes,
    SUM(CASE WHEN c.cartao = 'Vermelho' THEN 1 ELSE 0 END) AS vermelhos
FROM cartoes c
JOIN partidas p ON c.partida_id = p.partida_id
GROUP BY p.partida_id, p.data, p.mandante, p.gols_mandante, p.gols_visitante, p.visitante
ORDER BY total_cartoes DESC
LIMIT 10;""")

subtitulo("3.4 Window Functions (funções analíticas)")
descricao("Ranking de pontos dentro de cada temporada (top 4 e zona de rebaixamento):")
sql("""\
SELECT
    temporada, clube, pontos,
    RANK() OVER (PARTITION BY temporada ORDER BY pontos DESC) AS posicao,
    CASE
        WHEN RANK() OVER (PARTITION BY temporada ORDER BY pontos DESC) <= 4
             THEN 'Libertadores'
        WHEN RANK() OVER (PARTITION BY temporada ORDER BY pontos DESC) <= 6
             THEN 'Sul-Americana'
        WHEN RANK() OVER (PARTITION BY temporada ORDER BY pontos DESC) >= 17
             THEN 'Rebaixamento'
        ELSE ''
    END AS zona
FROM classificacao_historica
WHERE temporada = 2024
ORDER BY pontos DESC;""")

descricao("Diferença de pontos entre campeão e vice por temporada:")
sql("""\
WITH ranked AS (
    SELECT temporada, clube, pontos,
           RANK() OVER (PARTITION BY temporada ORDER BY pontos DESC) AS pos
    FROM classificacao_historica
),
campeao AS (SELECT temporada, clube, pontos FROM ranked WHERE pos = 1),
vice    AS (SELECT temporada, clube, pontos FROM ranked WHERE pos = 2)
SELECT c.temporada, c.clube AS campeao, c.pontos AS pts_campeao,
       v.clube AS vice, v.pontos AS pts_vice,
       (c.pontos - v.pontos) AS diferenca
FROM campeao c
JOIN vice v ON c.temporada = v.temporada
ORDER BY c.temporada DESC;""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NÍVEL 4 — ANÁLISES ESPECIAIS
# ══════════════════════════════════════════════════════════════════════════════
titulo("Nível 4 — Análises Especiais")

subtitulo("4.1 Rivalidades históricas")
descricao("Retrospecto completo entre Flamengo e Fluminense (Fla-Flu):")
sql("""\
SELECT
    COUNT(*) AS jogos,
    SUM(CASE WHEN mandante='Flamengo'   AND gols_mandante > gols_visitante THEN 1
             WHEN visitante='Flamengo'  AND gols_visitante > gols_mandante THEN 1
             ELSE 0 END) AS vitorias_fla,
    SUM(CASE WHEN gols_mandante = gols_visitante THEN 1 ELSE 0 END) AS empates,
    SUM(CASE WHEN mandante='Fluminense' AND gols_mandante > gols_visitante THEN 1
             WHEN visitante='Fluminense' AND gols_visitante > gols_mandante THEN 1
             ELSE 0 END) AS vitorias_flu,
    SUM(gols_mandante + gols_visitante) AS total_gols
FROM partidas
WHERE ((mandante = 'Flamengo'   AND visitante = 'Fluminense')
    OR (mandante = 'Fluminense' AND visitante = 'Flamengo'))
  AND gols_mandante IS NOT NULL;""")

descricao("Confronto direto entre dois times quaisquer (altere os nomes):")
sql("""\
SET @time1 = 'Palmeiras';
SET @time2 = 'Corinthians';

SELECT temporada, data, rodada, mandante, gols_mandante, gols_visitante, visitante
FROM partidas
WHERE ((mandante = @time1 AND visitante = @time2)
    OR (mandante = @time2 AND visitante = @time1))
  AND gols_mandante IS NOT NULL
ORDER BY data;""")

subtitulo("4.2 Maiores goleadas da história")
descricao("Top 10 goleadas com maior diferença de placar:")
sql("""\
SELECT
    temporada, data, mandante, gols_mandante, gols_visitante, visitante,
    ABS(gols_mandante - gols_visitante) AS diferenca
FROM partidas
WHERE gols_mandante IS NOT NULL
ORDER BY diferenca DESC, (gols_mandante + gols_visitante) DESC
LIMIT 10;""")

subtitulo("4.3 Sequências e consistência")
descricao("Clubes que mais vezes ficaram no G4 (top 4) ao longo da história:")
sql("""\
SELECT clube, COUNT(*) AS vezes_no_g4
FROM rebaixamento_acesso
WHERE posicao <= 4
GROUP BY clube
ORDER BY vezes_no_g4 DESC
LIMIT 10;""")

descricao("Temporadas com maior média de público (se disponível) ou mais gols por rodada:")
sql("""\
SELECT
    temporada,
    rodada,
    SUM(gols_mandante + gols_visitante) AS gols_rodada,
    COUNT(*) AS partidas
FROM partidas
WHERE gols_mandante IS NOT NULL
GROUP BY temporada, rodada
ORDER BY gols_rodada DESC
LIMIT 15;""")

subtitulo("4.4 Análise de 2026 (temporada em andamento)")
descricao("Situação atual de 2026 — apenas jogos finalizados:")
sql("""\
SELECT
    ROW_NUMBER() OVER (ORDER BY
        (3 * SUM(CASE WHEN gols_mandante > gols_visitante THEN 1
                      WHEN gols_visitante > gols_mandante THEN 0
                      ELSE 1 END) -- empate mandante
              + SUM(CASE WHEN gols_mandante = gols_visitante THEN 1 ELSE 0 END))
        DESC) AS pos,
    mandante AS clube,
    COUNT(*) AS jogos,
    SUM(CASE WHEN gols_mandante > gols_visitante THEN 1 ELSE 0 END) AS V,
    SUM(CASE WHEN gols_mandante = gols_visitante THEN 1 ELSE 0 END) AS E,
    SUM(CASE WHEN gols_mandante < gols_visitante THEN 1 ELSE 0 END) AS D
FROM partidas
WHERE temporada = 2026 AND status = 'finished'
GROUP BY mandante
ORDER BY V DESC, E DESC;""")

descricao("Forma recente — últimas 5 partidas de cada time em 2026:")
sql("""\
SELECT p.mandante AS clube, p.rodada, p.gols_mandante AS gf, p.gols_visitante AS gc,
       CASE WHEN p.gols_mandante > p.gols_visitante THEN 'V'
            WHEN p.gols_mandante = p.gols_visitante THEN 'E'
            ELSE 'D' END AS resultado
FROM partidas p
WHERE p.temporada = 2026 AND p.status = 'finished'
  AND p.mandante IN ('Palmeiras','Flamengo','Fluminense','Botafogo','São Paulo')
ORDER BY p.mandante, p.rodada DESC;""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DICAS FINAIS
# ══════════════════════════════════════════════════════════════════════════════
titulo("Dicas e Referência Rápida")

subtitulo("Filtros úteis para a coluna 'status'")
descricao("Use na cláusula WHERE para controlar quais jogos incluir:")
sql("""\
-- Apenas jogos já encerrados (use para análises históricas)
WHERE status = 'finished'

-- Jogos agendados / futuros (temporada 2026)
WHERE status = 'notStarted'

-- Sem filtro: inclui todos (histórico + agendados)""")

subtitulo("Nomes dos clubes mais usados")
sql("""\
SELECT DISTINCT mandante AS clube FROM partidas ORDER BY mandante;
-- Exemplos comuns: Flamengo, Palmeiras, Corinthians, São Paulo, Botafogo,
--   Fluminense, Internacional, Grêmio, Atlético Mineiro, Cruzeiro,
--   Athletico, Vasco da Gama, Santos, Bahia, Fortaleza""")

subtitulo("Template para análise de qualquer clube")
sql("""\
-- Substitua 'Palmeiras' pelo clube desejado
SELECT
    c.temporada,
    c.pontos,
    c.vitorias AS V, c.empates AS E, c.derrotas AS D,
    c.gols_pro AS GP, c.gols_contra AS GC,
    (c.gols_pro - c.gols_contra) AS SG,
    r.posicao AS classificacao_final
FROM classificacao_historica c
LEFT JOIN rebaixamento_acesso r
       ON c.temporada = r.temporada AND c.clube = r.clube
WHERE c.clube = 'Palmeiras'
ORDER BY c.temporada;""")

subtitulo("Como atualizar os dados de 2026")
descricao(
    "Os dados de 2026 são atualizados executando o pipeline Python novamente. "
    "No terminal, dentro da pasta do projeto:"
)
sql("""\
-- No terminal (cmd/PowerShell), dentro da pasta datalake/pipelines/:
python ingest_sofascore.py --seasons 2026 --fixtures-only
python build_gold.py
python load_mysql.py""")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Projeto Brasileirao Data Lake  |  Gerado automaticamente  |  2026")
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.size = Pt(9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Salva ──────────────────────────────────────────────────────────────────────
out = r"c:\Users\leeof\OneDrive\Documents\Documentos\geral downloadas\Brasileirao_Dataset-master\Brasileirao_Queries_Treinamento.docx"
doc.save(out)
print(f"Salvo em: {out}")
